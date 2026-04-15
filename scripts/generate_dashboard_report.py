"""
Generate a Word document explaining the Battery AI Co-Scientist Dashboard.
Run: python scripts/generate_dashboard_report.py
"""
from __future__ import annotations
from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT    = Path(__file__).resolve().parents[1]
IMG_DIR = ROOT / "outputs" / "report_imgs"
IMG_DIR.mkdir(parents=True, exist_ok=True)

NAVY  = RGBColor(0x1F, 0x3A, 0x5F)
TEAL  = RGBColor(0x00, 0x7A, 0x8A)
GREY  = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2E, 0x86, 0x48)
RED   = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xE6, 0x8A, 0x00)

def _shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Cm(widths[i])

def heading(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p

def body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    return p

def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = TEAL

# ── Diagram: two dashboards side by side ──────────────────────────────────────
def make_dashboard_overview():
    fig, axes = plt.subplots(1, 2, figsize=(13, 4))

    for ax, title, color, items in [
        (axes[0], "React Dashboard (BatteryDashboard.jsx)\nfor embedding in a web app", "#1F3A5F", [
            "Battery Selector (toggle on/off per battery)",
            "Tab: Comparison — RUL trajectory + capacity fade overlay + table",
            "Tab: Risk & Uncertainty — conformal coverage per temperature group",
            "Tab: Drift Monitor — PSI bar chart per feature",
            "Tab: Model QA — RMSE / MAE / shape error cards",
            "Tab: Hostile Tests — adversarial test pass/fail table",
            "Live header: XGB RMSE, TCN RMSE, Coverage, Drift alert",
        ]),
        (axes[1], "Streamlit Dashboard (app.py)\nfor running locally with Python", "#007A8A", [
            "Artifact status panel (which pipeline outputs exist)",
            "Multi-battery comparison: capacity overlay + RUL overlay",
            "RUL uncertainty bands (5th-95th percentile toggle)",
            "Anomaly markers overlaid on capacity curve",
            "Single-battery: degradation + anomalies + uncertainty + survival",
            "Hazard probability and horizon failure probability charts",
            "Final system report rendered inline",
        ]),
    ]:
        ax.set_facecolor("#F9F9F9")
        ax.axis("off")
        ax.text(0.5, 0.97, title, ha="center", va="top", fontsize=9,
                fontweight="bold", color=color, transform=ax.transAxes)
        for i, item in enumerate(items):
            y = 0.82 - i * 0.115
            ax.add_patch(mpatches.FancyBboxPatch(
                (0.03, y - 0.04), 0.94, 0.09,
                boxstyle="round,pad=0.02", facecolor=color, alpha=0.12,
                edgecolor=color, linewidth=0.6, transform=ax.transAxes))
            ax.text(0.06, y + 0.005, f"  {item}", ha="left", va="center",
                    fontsize=7.5, color="#2C2C2C", transform=ax.transAxes)

    fig.patch.set_facecolor("#F9F9F9")
    plt.tight_layout(pad=0.5)
    path = IMG_DIR / "dashboard_overview.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def make_tab_diagram():
    fig, ax = plt.subplots(figsize=(12, 2.5))
    ax.axis("off")
    ax.set_facecolor("#0F172A")
    fig.patch.set_facecolor("#0F172A")

    tabs = [
        ("Comparison", "#F59E0B", True),
        ("Risk & Uncertainty", "#94A3B8", False),
        ("Drift Monitor", "#94A3B8", False),
        ("Model QA", "#94A3B8", False),
        ("Hostile Tests", "#94A3B8", False),
    ]
    x = 0.02
    for label, color, active in tabs:
        w = len(label) * 0.013 + 0.04
        bg = "#1E293B" if active else "#0F172A"
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, 0.3), w, 0.5, boxstyle="round,pad=0.02",
            facecolor=bg, edgecolor=color if active else "#334155", linewidth=1.5,
            transform=ax.transAxes))
        ax.text(x + w/2, 0.57, label, ha="center", va="center",
                fontsize=8, color=color, fontweight="bold" if active else "normal",
                transform=ax.transAxes, fontfamily="monospace")
        x += w + 0.015

    ax.text(0.5, 0.05, "Active tab shown in amber.  Click any tab to switch views.",
            ha="center", va="bottom", fontsize=8, color="#64748B",
            transform=ax.transAxes, fontfamily="monospace")

    path = IMG_DIR / "tabs_diagram.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def make_data_flow():
    fig, ax = plt.subplots(figsize=(12, 2.8))
    ax.axis("off"); ax.set_facecolor("#F9F9F9"); fig.patch.set_facecolor("#F9F9F9")

    boxes = [
        ("Pipeline runs\n(Stages 1–6)", "#1F3A5F"),
        ("JSON / CSV\noutput files", "#4A4A4A"),
        ("Dashboard\nloads files", "#007A8A"),
        ("Charts &\ntables render", "#2E8648"),
        ("You read\nthe results", "#5B4A8A"),
    ]
    x = 0.04
    for label, color in boxes:
        w = 0.15
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, 0.2), w, 0.55, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor="white", linewidth=1.5, transform=ax.transAxes))
        ax.text(x + w/2, 0.48, label, ha="center", va="center",
                fontsize=8, color="white", fontweight="bold", transform=ax.transAxes)
        if label != boxes[-1][0]:
            ax.annotate("", xy=(x + w + 0.055, 0.47), xytext=(x + w + 0.005, 0.47),
                        xycoords="axes fraction", textcoords="axes fraction",
                        arrowprops=dict(arrowstyle="->", color="#AAAAAA", lw=1.5))
        x += w + 0.065

    path = IMG_DIR / "data_flow.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

# ── BUILD ─────────────────────────────────────────────────────────────────────
def build():
    print("Generating figures...")
    img_overview = make_dashboard_overview()
    img_tabs     = make_tab_diagram()
    img_flow     = make_data_flow()
    print("Building document...")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # Title
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Battery Dashboard — Explained")
    tr.font.size = Pt(26); tr.font.bold = True; tr.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("A plain-English guide to what the dashboard shows,\nhow it works, and where the data comes from")
    sr.font.size = Pt(13); sr.font.italic = True; sr.font.color.rgb = TEAL
    doc.add_page_break()

    # ── 1. What is the dashboard ──────────────────────────────────────────────
    heading(doc, "1. What Is the Dashboard?")
    body(doc,
        "The dashboard is a visual interface that reads the output files produced by the "
        "pipeline and displays them as charts and tables. It does not run any models — "
        "it only shows results that have already been computed. Think of it like a "
        "report viewer, not a calculator.")
    body(doc,
        "There are actually two separate dashboards in this project:")
    bullet(doc, "React Dashboard (BatteryDashboard.jsx) — a modern web component written in React/JavaScript. "
                "Designed to be embedded inside a web application. Has 5 tabs. Dark-mode design.")
    bullet(doc, "Streamlit Dashboard (app.py) — a Python-based dashboard you run from the terminal. "
                "Opens in your browser. Better for quick exploration and debugging. Shows more raw detail.")
    body(doc,
        "Both dashboards read the exact same pipeline output files. They just present "
        "them differently.")
    doc.add_paragraph()

    doc.add_picture(str(img_overview), width=Inches(6.5))
    caption(doc, "Figure 1 — Side-by-side comparison of what each dashboard shows.")

    # ── 2. How does the data get there ────────────────────────────────────────
    heading(doc, "2. How Does the Dashboard Get Its Data?")
    body(doc,
        "The dashboard does not connect to a live sensor or a database. "
        "It reads static files that the pipeline saved after running. "
        "Here is the flow:")
    doc.add_picture(str(img_flow), width=Inches(6.2))
    caption(doc, "Figure 2 — Data flows from pipeline output files into the dashboard. No live computation happens.")

    body(doc, "The specific files the dashboard reads:")
    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["File", "What is in it", "Used by which tab"]):
        t.rows[0].cells[j].text = h
        _shade(t.rows[0].cells[j], "1F3A5F")
        for para in t.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE; run.font.bold = True; run.font.size = Pt(9)
    rows = [
        ("uncertainty_estimates.json", "RUL prediction + lower/upper bounds + risk category for every cycle of every battery", "Comparison, Risk & Uncertainty"),
        ("conformal_coverage_report.json", "How well the 90% coverage target was met, split by temperature group", "Risk & Uncertainty"),
        ("drift_report.json", "PSI score for each feature — how different the test data is from training data", "Drift Monitor"),
        ("model_metrics.json", "RMSE, MAE for XGBoost, TCN, statistical baseline. Also CV results.", "Model QA, header bar"),
        ("cycle_features_with_rul.csv", "The cleaned cycle-level data (capacity, temperature, RUL, etc.)", "Comparison (capacity chart)"),
        ("hostile_results.json", "Results of adversarial/edge-case robustness tests", "Hostile Tests"),
    ]
    for i, (f, w, u) in enumerate(rows, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = w
        t.rows[i].cells[2].text = u
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _col_widths(t, [4.5, 7.5, 4])
    doc.add_paragraph()

    # ── 3. React dashboard tabs ───────────────────────────────────────────────
    heading(doc, "3. React Dashboard — Tab by Tab")
    body(doc,
        "The React dashboard has 5 tabs. At the top there is always a header bar "
        "showing the three most important numbers at a glance: XGBoost RMSE, TCN RMSE, "
        "and the overall conformal coverage percentage.")
    doc.add_picture(str(img_tabs), width=Inches(6.0))
    caption(doc, "Figure 3 — The five tabs of the React dashboard.")

    # Tab 1
    heading(doc, "Tab 1: Comparison", 2)
    body(doc,
        "This is the default tab. It lets you select which batteries to look at and "
        "compare them side by side.")
    body(doc, "What you see:", bold=True)
    bullet(doc, "Battery Selector — pill buttons at the top, colour-coded by temperature group "
                "(green = room, purple = hot, blue = cold). Click to toggle a battery on or off.")
    bullet(doc, "Chart toggle — switch between 'RUL Trajectory' (how many cycles are left over time) "
                "and 'Capacity Fade' (how the actual capacity is dropping over cycles).")
    bullet(doc, "Line chart — one line per selected battery, all drawn on the same axes so you can "
                "compare degradation speeds directly.")
    bullet(doc, "Comparison table — one row per battery showing: number of cycles observed, "
                "final predicted RUL, risk badge (HIGH/MEDIUM/LOW), failure probability, "
                "latest capacity reading, conformal interval width, q_hat, and empirical coverage.")
    note(doc, "Tip: A battery with a large 'Interval Width' means the model is less certain about "
              "its RUL — the true value could be anywhere in a wide range.")

    # Tab 2
    heading(doc, "Tab 2: Risk and Uncertainty", 2)
    body(doc,
        "This tab shows whether the confidence intervals (the range around each RUL prediction) "
        "are actually hitting the 90% target.")
    body(doc, "What you see:", bold=True)
    bullet(doc, "Three cards — one per temperature group (room, hot, cold). Each card shows the "
                "empirical coverage achieved for that group, the q_hat value used, and the calibration strategy.")
    bullet(doc, "If the number is red, that group is under-covered (below the 90% target).")
    bullet(doc, "Strategy shown is either 'split' (standard conformal) or 'lobo' (leave-one-battery-out, "
                "used for cold batteries to improve coverage).")
    note(doc, "In our results: room = 100%, hot = 100%, cold = 88.7% (slightly below target — "
              "due to limited cold-battery calibration data).")

    # Tab 3
    heading(doc, "Tab 3: Drift Monitor", 2)
    body(doc,
        "This tab answers: 'is the test data coming from the same distribution as training data?' "
        "If the distributions are very different, model predictions may be less reliable.")
    body(doc, "What you see:", bold=True)
    bullet(doc, "Four summary cards: overall status, features monitored, active alerts, "
                "and degradation-related flags.")
    bullet(doc, "Horizontal bar chart — one bar per feature, coloured by PSI severity: "
                "green (no drift), amber (slight drift), red (significant drift).")
    bullet(doc, "Two dashed reference lines: amber threshold at PSI=0.10, red threshold at PSI=0.20.")
    bullet(doc, "Degradation features (capacity, ah_est, energy_j) are shown in grey — their drift "
                "is expected because capacity decreases by design.")
    note(doc, "Why do we have RED alerts? Because test batteries include cold-temperature batteries "
              "that were not well represented in training. Temperature affects everything else "
              "(voltage, current, duration), so high PSI on those features is expected and explained.")

    # Tab 4
    heading(doc, "Tab 4: Model QA", 2)
    body(doc,
        "A quick quality check panel showing the core accuracy metrics for all three models.")
    body(doc, "What you see:", bold=True)
    bullet(doc, "XGBoost RMSE card — 22.08 cycles (primary model, best performer)")
    bullet(doc, "TCN RMSE card — 25.63 cycles (deep learning sequence model)")
    bullet(doc, "Baseline RMSE card — 466.82 cycles (exponential decay, shown in red as a reference)")
    bullet(doc, "Shape Error — 0.311 (measures how well the prediction curve shape matches the "
                "true RUL trajectory; 0 = perfect, 1 = completely wrong shape)")
    note(doc, "RMSE = Root Mean Squared Error. Lower is better. It is measured in cycles, so "
              "RMSE=22 means the model is off by about 22 cycles on average.")

    # Tab 5
    heading(doc, "Tab 5: Hostile Tests", 2)
    body(doc,
        "This tab shows results from the adversarial robustness tests — deliberate bad inputs "
        "fed to the pipeline to check it handles them safely.")
    body(doc, "What you see:", bold=True)
    bullet(doc, "Summary cards: total tests, passed, failed, pass rate.")
    bullet(doc, "Test result grid — each row is one test case with its ID, description, "
                "PASS/FAIL result, and the confidence score the system gave.")
    bullet(doc, "A test 'passes' if the system either rejects the bad input with a warning "
                "or returns a low confidence score (not a confident wrong answer).")
    note(doc, "Examples of hostile tests: completely missing columns, voltages outside physical "
              "range, all-zero current readings, truncated time series with only 2 cycles.")

    doc.add_page_break()

    # ── 4. Streamlit dashboard ────────────────────────────────────────────────
    heading(doc, "4. Streamlit Dashboard (app.py) — Section by Section")
    body(doc,
        "The Streamlit dashboard is run from the terminal. You open a terminal, "
        "navigate to the project folder, and run:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run("streamlit run dashboard/app.py")
    r.font.name = "Courier New"; r.font.size = Pt(11); r.font.color.rgb = NAVY; r.font.bold = True
    body(doc, "Your browser then opens automatically showing the dashboard.")
    doc.add_paragraph()

    sections = [
        ("Artifact Status Panel",
         "The very first thing you see. Shows which output files exist and "
         "four summary numbers: total cycle rows, anomalies count, uncertainty rows, "
         "survival rows. If a file is missing it shows False in the table — telling you "
         "which pipeline stage failed or hasn't been run yet."),
        ("Battery Selector",
         "A dropdown to pick one battery for the single-battery detailed views "
         "(degradation chart, uncertainty chart, survival chart). Separate from the "
         "multi-battery comparison below."),
        ("Battery Comparison (Multi-select)",
         "A multi-select box where you can choose any combination of batteries. "
         "Shows a capacity degradation overlay (all selected batteries on one chart, "
         "with red dots marking anomaly cycles) and an RUL overlay. "
         "You can toggle uncertainty bands (the shaded 5th-95th percentile region) "
         "on/off. Also shows a comparison table with RUL, risk category, failure "
         "probability, and interval width for the latest cycle of each battery."),
        ("Degradation Trajectory + Anomalies",
         "Single-battery view. Plots capacity over cycles for the selected battery. "
         "Red dots are cycles flagged as anomalies by the anomaly detector."),
        ("Uncertainty",
         "Single-battery view. Plots predicted RUL over cycles with a shaded "
         "blue band showing the 90% conformal interval. Wider band = more uncertainty. "
         "The band should narrow as the battery reaches end-of-life and the model "
         "becomes more certain about the remaining cycles."),
        ("Survival / Hazard Risk",
         "Single-battery view. Two stacked charts: top shows the per-cycle hazard "
         "probability (chance of failing at this exact cycle), bottom shows the "
         "20-cycle horizon failure probability (chance of failing within the next "
         "20 cycles from this point). Risk category (HIGH/MEDIUM/LOW) comes from "
         "the horizon probability using the 70%/30% thresholds."),
        ("Final System Report",
         "The supervisor report from Stage 6 is rendered inline at the bottom. "
         "You can read the full CONDITIONAL PASS verdict with all stage-by-stage "
         "checks without leaving the dashboard."),
    ]

    for title, desc in sections:
        heading(doc, title, 2)
        body(doc, desc)

    doc.add_page_break()

    # ── 5. How to read the numbers ────────────────────────────────────────────
    heading(doc, "5. How to Read the Numbers — Quick Reference")
    body(doc,
        "When you are presenting or explaining the dashboard, here is a plain-English "
        "guide to what each number means:")

    t2 = doc.add_table(rows=13, cols=3)
    t2.style = "Table Grid"
    for j, h in enumerate(["Number / term", "What it is", "Good value"]):
        t2.rows[0].cells[j].text = h
        _shade(t2.rows[0].cells[j], "1F3A5F")
        for para in t2.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE; run.font.bold = True; run.font.size = Pt(9)
    guide = [
        ("RUL prediction", "How many cycles the model thinks the battery has left", "As accurate as possible — ours is ~22 cycles off on average"),
        ("Interval width", "The width of the confidence band around the RUL prediction", "Smaller is more certain. Ours is ~83 cycles wide on average"),
        ("Coverage 98.3%", "In 98.3% of test observations, the true RUL fell inside the predicted interval", "Must be >= 90%. We exceed the target."),
        ("q_hat", "The margin added/subtracted to make the interval (e.g. q_hat=54 means interval is prediction +/- 54)", "Lower means the model needs less margin to stay accurate"),
        ("RMSE 22.08", "On average the prediction is ~22 cycles away from the true value", "Lower is better. Our 100-cycle threshold is the pass criterion."),
        ("Failure probability", "P(battery fails within the next 20 cycles)", "0 = definitely fine, 1 = definitely failing soon"),
        ("Risk: HIGH", "Failure probability > 70%", "Battery needs attention now"),
        ("Risk: MEDIUM", "Failure probability 30-70%", "Monitor and plan maintenance"),
        ("Risk: LOW", "Failure probability < 30%", "Normal operation"),
        ("PSI (drift)", "How different the test-data distribution is from training", "< 0.10 = fine, 0.10-0.20 = watch, > 0.20 = significant shift"),
        ("Hazard probability", "Chance of failing at this specific cycle given it survived this far", "Low early in life, rises steeply near end of life"),
        ("Shape error 0.311", "How well the prediction curve follows the true RUL curve (1-correlation)", "0 = perfect, 1 = no shape match"),
    ]
    for i, (n, w, g) in enumerate(guide, 1):
        t2.rows[i].cells[0].text = n
        t2.rows[i].cells[1].text = w
        t2.rows[i].cells[2].text = g
        for cell in t2.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _col_widths(t2, [4, 6.5, 5.5])
    doc.add_paragraph()

    # ── 6. Why two dashboards ─────────────────────────────────────────────────
    heading(doc, "6. Why Are There Two Dashboards?")
    t3 = doc.add_table(rows=7, cols=3)
    t3.style = "Table Grid"
    for j, h in enumerate(["", "React (BatteryDashboard.jsx)", "Streamlit (app.py)"]):
        t3.rows[0].cells[j].text = h
        _shade(t3.rows[0].cells[j], "4A4A4A")
        for para in t3.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE; run.font.bold = True; run.font.size = Pt(9)
    compare = [
        ("Language",   "JavaScript (React)",        "Python (Streamlit)"),
        ("How to run", "Needs a React app / Node.js", "streamlit run app.py in terminal"),
        ("Best for",   "Deployment in a web app",   "Local exploration and debugging"),
        ("Charts",     "Recharts (interactive)",    "Matplotlib (static)"),
        ("Style",      "Dark mode, monospace font", "Standard Streamlit theme"),
        ("Extra features", "Hostile test tab, conformal tab, live header bar", "More raw data, survival charts, inline report"),
    ]
    for i, (label, react, streamlit) in enumerate(compare, 1):
        t3.rows[i].cells[0].text = label
        t3.rows[i].cells[1].text = react
        t3.rows[i].cells[2].text = streamlit
        for cell in t3.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _col_widths(t3, [3.5, 7, 5.5])

    out = ROOT / "Dashboard_Report.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    build()
