"""
Generate a Q&A Word document — simple and easy to understand for a first-time learner.
Run: python scripts/generate_qa_document.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]

NAVY  = RGBColor(0x1F, 0x3A, 0x5F)
TEAL  = RGBColor(0x00, 0x7A, 0x8A)
GREY  = RGBColor(0x3A, 0x3A, 0x3A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1E, 0x6E, 0x3A)
AMBER = RGBColor(0xAA, 0x60, 0x00)
LGREY = RGBColor(0xF5, 0xF5, 0xF5)

def _shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


QA_DATA = [
    # ── BASICS ──────────────────────────────────────────────────────────────
    ("BASICS", None, None),

    ("What is this project about, in one sentence?",
     "This project builds a computer system that looks at a lithium-ion battery and "
     "predicts how many more charge-discharge cycles it can do before it stops working properly.",
     None),

    ("What is a lithium-ion battery?",
     "It is the type of rechargeable battery used in phones, laptops, and electric cars. "
     "Every time you charge and discharge it, it gets very slightly weaker. "
     "Over hundreds of cycles, the battery can hold less and less charge.",
     None),

    ("What is a cycle?",
     "One cycle = one full charge plus one full discharge. "
     "So if you charge your phone from 0% to 100% and then use it until it dies, that is one cycle. "
     "The batteries in this project go through 100 to 500+ cycles before they reach end-of-life.",
     None),

    ("What is End-of-Life (EOL)?",
     "A battery has reached its End-of-Life when its capacity has dropped to 80% of what it was "
     "when it was brand new. For our batteries, they started at 2.0 Ah capacity, so EOL is when "
     "they reach 1.6 Ah. After this point the battery is considered worn out.",
     "EOL capacity = 80% x 2.0 Ah = 1.6 Ah"),

    ("What is RUL — Remaining Useful Life?",
     "RUL is the number of charge-discharge cycles a battery still has left before it hits "
     "End-of-Life. For example, if a battery currently has 150 cycles recorded and will reach "
     "EOL at cycle 300, its RUL right now is 150 cycles. The system's main job is to predict "
     "this number without knowing when EOL will actually happen.",
     None),

    ("Why is predicting RUL useful?",
     "If you know a battery is going to fail in 20 cycles, you can replace it before it breaks "
     "something important. In electric vehicles or medical devices, an unexpected battery failure "
     "can be dangerous or very expensive. RUL prediction enables preventive maintenance.",
     None),

    ("What dataset did we use?",
     "The NASA Prognostics Center of Excellence Battery Dataset. It contains real experimental "
     "data from 34 lithium-ion 18650 cylindrical cells (the same size as AA batteries but fatter). "
     "Each battery was repeatedly charged and discharged until it died, while voltage, current, "
     "and temperature were recorded every second.",
     None),

    ("What does 18650 mean?",
     "It is just the size code for that type of cylindrical battery cell. "
     "18 mm wide, 65 mm tall. These are the same cells used in Tesla cars and laptop batteries.",
     None),

    ("What are the three temperature groups?",
     "Room temperature (~24 degrees C): batteries B0005-B0036. Normal conditions.\n"
     "Hot (43 degrees C): batteries B0029-B0040. Accelerated aging.\n"
     "Cold (below 10 degrees C): batteries B0041-B0056. Slower but irregular degradation.\n"
     "Temperature matters a lot — hot batteries degrade faster, cold ones are harder to predict.",
     None),

    # ── THE PIPELINE ─────────────────────────────────────────────────────────
    ("THE PIPELINE / STAGES", None, None),

    ("What is a pipeline?",
     "In software, a pipeline is a series of steps that run one after another, where the output "
     "of each step becomes the input for the next. Our pipeline has 6 stages — like an assembly line "
     "for turning raw battery sensor data into a final report and predictions.",
     None),

    ("What does Stage 1-2 (Preprocessing) do?",
     "It takes the raw sensor data files (thousands of rows of voltage, current, temperature "
     "readings) and cleans them up. It: standardises column names (different files might call "
     "voltage 'V' or 'Voltage_measured'), checks all values are physically plausible, "
     "computes one summary row per cycle (instead of one row per second), and calculates "
     "the RUL label for each cycle. The output is one clean CSV file.",
     None),

    ("What does Stage 3 (Modeling) do?",
     "It trains three different mathematical models to predict RUL using the clean data. "
     "Think of it like three different experts each making their own prediction. "
     "The models are: an exponential decay formula (physics), XGBoost (machine learning), "
     "and a TCN neural network (deep learning).",
     None),

    ("What does Stage 4 (Uncertainty) do?",
     "Instead of just saying 'the RUL is 50 cycles', it adds a confidence range: "
     "'the RUL is 50 cycles, and I am 90% confident the true value is between 30 and 70 cycles'. "
     "This uses a technique called conformal prediction.",
     None),

    ("What does Stage 4.2 (Survival Risk) do?",
     "It calculates the probability that a battery will fail within the next 20 cycles. "
     "This is different from predicting an exact RUL number — it is asking 'how urgently "
     "should I be worried?' It uses a technique called Kaplan-Meier survival analysis.",
     None),

    ("What does Stage 4.5 (Anomaly Detection) do?",
     "It flags any cycles where the battery behaved unexpectedly — for example, a sudden "
     "capacity drop that was much larger than the gradual trend. These are outlier events "
     "worth investigating. 32 anomalies were found in our test batteries.",
     None),

    ("What does Stage 5 (Reasoning) do?",
     "It generates written hypotheses explaining WHY a battery might be degrading the way it is. "
     "For example: 'high mean temperature is associated with faster capacity loss in this battery.' "
     "It also creates counterfactual examples: 'if the energy per cycle had been 10% lower, "
     "the predicted RUL would be 15 cycles higher.' These are always labelled as hypotheses, "
     "not proven facts.",
     None),

    ("What does Stage 6 (Supervisor Review) do?",
     "It is the final quality check. An automated auditor goes through all the outputs and "
     "decides whether the system passes. It checks: Is the RMSE below the threshold? "
     "Is the conformal coverage report present? Were hypotheses and counterfactuals generated? "
     "It then writes a final verdict: PASS, CONDITIONAL PASS, or FAIL.",
     None),

    # ── THE MODELS ───────────────────────────────────────────────────────────
    ("THE MODELS", None, None),

    ("Why did we use three models instead of just one?",
     "Each model captures something different. The exponential model uses physics knowledge. "
     "XGBoost finds complex patterns in many features at once. The TCN sees sequences of "
     "cycles, not just one snapshot. By comparing them and combining the two best ones, "
     "we get more reliable predictions than any single model alone.",
     None),

    ("What is XGBoost in simple terms?",
     "XGBoost is a machine learning method that builds a 'forest' of simple decision trees "
     "one at a time. Each new tree focuses on fixing the mistakes of all the previous trees. "
     "The final prediction is the combined vote of all 400 trees. It is very widely used in "
     "industry because it works well on tabular data (rows and columns) without needing huge "
     "amounts of data.",
     None),

    ("What is a TCN (Temporal Convolutional Network)?",
     "It is a type of neural network designed for sequences of data over time. "
     "Instead of looking at just the current cycle, it looks at the last 8 cycles together "
     "and finds patterns in how the battery has been behaving recently. It uses special "
     "layers called dilated convolutions that let it look back farther with fewer calculations.",
     None),

    ("Why did the statistical baseline model fail so badly (RMSE 467)?",
     "The simple exponential decay formula assumes every battery degrades in a smooth, "
     "perfectly exponential curve. Real batteries do not — temperature changes, different "
     "charging rates, and manufacturing variations all cause deviations. The formula has "
     "no way to account for temperature or current history, so it makes large errors for "
     "batteries that behave differently from the average curve.",
     None),

    ("Why was the statistical model still included?",
     "As a physics-grounded reference point. If the ML model ever starts making predictions "
     "that are wildly inconsistent with the physics model, that is a warning sign. "
     "It also helps explain the degradation mechanism to non-technical audiences.",
     None),

    ("What does RMSE mean?",
     "RMSE stands for Root Mean Squared Error. It measures, on average, how far the "
     "predictions are from the true values. An RMSE of 22 cycles means the model's prediction "
     "is typically off by about 22 cycles. Lower RMSE = more accurate model.",
     "Formula: RMSE = square root of (average of all (prediction - true value) squared)"),

    ("What does MAE mean?",
     "MAE stands for Mean Absolute Error. It is the average of |prediction - true value| "
     "for all test examples. Less sensitive to large outlier errors than RMSE. "
     "Our XGBoost MAE is 13.18 cycles.",
     None),

    ("What is cross-validation and why does it matter?",
     "Cross-validation is a way to test your model more thoroughly than a single train/test split. "
     "We split the 22 training batteries into 5 groups. We train on 4 groups and test on 1, "
     "then repeat 5 times with a different group as the test each time. "
     "The average RMSE across all 5 tests (48.17) is a more honest accuracy estimate than "
     "the single holdout test result (22.08).",
     None),

    ("Why is the cross-validation RMSE (48) so much higher than the test RMSE (22)?",
     "The 6 test batteries we happened to use for the final evaluation turned out to be "
     "easier to predict than average. Cross-validation averages over all possible battery "
     "combinations, giving a more realistic picture. This is why the system gets CONDITIONAL PASS "
     "rather than a full PASS — it is an honest acknowledgement, not a failure.",
     None),

    ("What does 'battery-level split' mean and why does it matter?",
     "It means when we decide which batteries go into training and which into testing, "
     "we never split a battery between both. If battery B0030 is in the test set, ALL "
     "its cycles are in the test set. This is important because a model that has seen "
     "some cycles from B0030 would already know its degradation pattern — an unfair advantage "
     "that makes accuracy look better than it really is. This is called data leakage.",
     None),

    # ── UNCERTAINTY / CONFORMAL PREDICTION ───────────────────────────────────
    ("UNCERTAINTY AND CONFIDENCE INTERVALS", None, None),

    ("Why do we need a confidence interval? Why not just give one number?",
     "A single number is misleading. If the model says 'RUL is 50 cycles' but the true "
     "value could be anywhere from 10 to 90, giving a single number creates false confidence. "
     "A confidence interval says 'I am 90% sure the true value is between 30 and 70 cycles' — "
     "much more honest and useful for decision-making.",
     None),

    ("What is conformal prediction?",
     "It is a method that wraps around any model and adds statistically guaranteed confidence "
     "intervals. The guarantee is: if you follow this procedure, the true value will fall inside "
     "the interval at least 90% of the time, regardless of which model you use underneath. "
     "The key word is 'guaranteed' — unlike standard model confidence estimates, this actually holds.",
     None),

    ("How does conformal prediction work — the simple version?",
     "Step 1: Run the model on a set of batteries it never trained on (the calibration set). "
     "Compute the error on each: |true RUL - predicted RUL|.\n"
     "Step 2: Find the 90th percentile of those errors. Call it q_hat.\n"
     "Step 3: For any new battery, the interval is simply: [prediction - q_hat, prediction + q_hat].\n"
     "Because q_hat was chosen to cover 90% of past errors, it covers 90% of future errors too.",
     "q_hat = the margin. Bigger margin = wider interval = more certain to contain the true value."),

    ("What is LOBO and why was it needed?",
     "LOBO stands for Leave-One-Battery-Out. The problem was that cold batteries were getting "
     "a lower-than-target coverage (87.3% instead of 90%). The reason: we only had 3 cold "
     "calibration batteries, so the errors on those 3 might not represent all cold batteries well.\n\n"
     "LOBO fixes this by: for each cold battery, temporarily removing it, retraining the model "
     "without it, and measuring the error. This gives errors that are more representative of "
     "what happens on a truly unseen cold battery. Result: coverage improved to 88.7%.",
     None),

    ("What does 'coverage 98.3%' mean?",
     "It means that in 98.3% of test observations, the true RUL value fell inside our predicted "
     "interval. Our target was 90%. We exceeded it — the intervals are a bit wider than necessary "
     "but that is safer than being too narrow.",
     None),

    ("Why is cold battery coverage only 88.7%, slightly below the 90% target?",
     "We only had 3 cold batteries in the calibration set. With so few examples, it is hard to "
     "perfectly calibrate the q_hat. More cold-battery data would fix this. The gap is only "
     "1.3 percentage points and is clearly explained — it is not a random failure.",
     None),

    # ── SURVIVAL ANALYSIS ────────────────────────────────────────────────────
    ("SURVIVAL ANALYSIS", None, None),

    ("What is survival analysis?",
     "It is a branch of statistics originally developed for medical research to answer: "
     "'what is the probability that a patient (or in our case, a battery) survives past a "
     "certain time point?' It handles both batteries that have failed AND batteries that have "
     "not failed yet (censored observations).",
     None),

    ("What does 'censored' mean in this context?",
     "A battery is censored if it had not yet reached End-of-Life by the last cycle we have "
     "data for. We know it survived at least that long, but we don't know exactly when it "
     "would have failed. Survival analysis correctly accounts for this partial information "
     "instead of throwing it away.",
     None),

    ("What is the Kaplan-Meier method?",
     "It is the simplest and most common survival analysis technique. It estimates the "
     "probability of surviving past any cycle t using only the observed failure times. "
     "No assumptions about the shape of the distribution are needed. "
     "It is called non-parametric for this reason.",
     "Formula: S(t) = multiply (1 - failures_at_t / at_risk_at_t) for all t up to now"),

    ("What is the hazard rate?",
     "The hazard rate h(t) is the probability that the battery fails at this specific cycle, "
     "given that it has survived up to this point. It is calculated from the survival curve "
     "as: h(t) = 1 - S(t+1)/S(t). Early in a battery's life the hazard is very low. "
     "Near end-of-life it rises steeply.",
     None),

    ("What is the 'horizon failure probability' and why 20 cycles?",
     "The horizon failure probability answers: given we are at cycle t right now, what is the "
     "probability the battery fails within the NEXT 20 cycles? "
     "Formula: P(fail within 20) = 1 - S(t + 20) / S(t).\n\n"
     "The 20-cycle horizon is a configurable setting in pipeline.yaml. It represents a "
     "maintenance planning window — if failure probability is high within 20 cycles, "
     "you schedule maintenance in the next 20 cycles.",
     None),

    ("How are the risk thresholds (70% and 30%) decided?",
     "They are engineering thresholds set in the configuration file (configs/pipeline.yaml). "
     "They represent a judgment call: above 70% failure probability in the next 20 cycles "
     "is considered HIGH risk (requires immediate action), between 30-70% is MEDIUM (monitor "
     "carefully), below 30% is LOW (normal operation). These are not derived from a formula — "
     "they are the project team's chosen operating thresholds. In a real deployment they would "
     "be agreed with the engineering team responsible for the batteries.",
     None),

    # ── DRIFT / PSI ───────────────────────────────────────────────────────────
    ("DRIFT MONITORING", None, None),

    ("What is feature drift?",
     "After a model is trained, the new data it sees in the real world might be different from "
     "the data it was trained on. For example, if the model was trained mostly on room-temperature "
     "batteries but is now being asked to predict cold batteries, the temperature distribution has "
     "drifted. This can make predictions less reliable.",
     None),

    ("What is PSI?",
     "PSI stands for Population Stability Index. It measures how different the distribution of "
     "a feature is between training data and new data. A low PSI means the distributions are "
     "similar. A high PSI means they have drifted.",
     "PSI thresholds: < 0.10 = fine (GREEN), 0.10-0.20 = watch (AMBER), > 0.20 = problem (RED)"),

    ("Why are so many features showing RED drift in our results?",
     "Because our test set includes 3 cold batteries (B0041, B0044, B0052) which were tested "
     "at a much lower temperature than most training batteries. Lower temperature changes "
     "everything: voltage behaviour, current, cycle duration. So the distributions of those "
     "features look very different between training (mostly room temperature) and test "
     "(mixed with cold). This is expected and explained, not a failure of the pipeline.",
     None),

    # ── ANOMALY DETECTION ────────────────────────────────────────────────────
    ("ANOMALY DETECTION", None, None),

    ("What is an anomaly in battery data?",
     "An anomaly is a cycle where the battery behaved significantly differently from what "
     "was expected given its degradation trend. For example: a sudden large capacity drop "
     "in one cycle followed by recovery, or a cycle where the prediction error was unusually large. "
     "These can indicate measurement problems, unusual operating conditions, or the start of "
     "accelerated degradation.",
     None),

    ("How are anomalies detected?",
     "The anomaly detector compares the actual RUL at each cycle to what the model predicted. "
     "If the absolute difference (|true - predicted|) is very large compared to the typical "
     "error for that battery, the cycle is flagged as an anomaly. Additionally, batteries "
     "B0049-B0056 had a pre-training anomaly filter applied because their capacities sometimes "
     "dropped below 50% of their median — likely measurement artefacts.",
     None),

    ("We found 32 anomalies. Is that bad?",
     "Not necessarily. 32 anomalies across 638 test cycles is about 5%. Zero anomalies "
     "would mean every cycle degraded perfectly smoothly, which is actually suspicious. "
     "Real batteries have occasional anomalous cycles. The important thing is that the "
     "anomalies are flagged so a human engineer can investigate them.",
     None),

    # ── VALIDATION AND ROBUSTNESS ─────────────────────────────────────────────
    ("VALIDATION AND ROBUSTNESS", None, None),

    ("What is the hostile validation suite?",
     "It is a set of deliberately bad inputs designed to stress-test the pipeline. "
     "For example: a file with completely missing columns, voltage readings of 999V "
     "(physically impossible), only 2 cycles of data (not enough to fit a model), "
     "all-zero current readings. The test passes if the system either rejects the input "
     "gracefully or returns a very low confidence score instead of making confident wrong predictions.",
     None),

    ("What is schema validation / validation gating?",
     "Before any modeling happens, each feature is checked against physical reality. "
     "Voltage must be between 2.0V and 4.2V. Temperature must be between 4C and 44C. "
     "Capacity must be between 1.0 Ah and 2.0 Ah. If values fall outside these ranges, "
     "the pipeline flags them and either rejects or warns. This prevents garbage data from "
     "producing confident but meaningless predictions.",
     None),

    ("What is schema mapping / semantic mapping?",
     "Different labs save data with different column names. One file might call voltage "
     "'Voltage_measured', another 'V', another 'volt'. The schema mapper automatically "
     "detects which columns correspond to which physical quantities using pattern matching "
     "and known aliases, then renames them to the standard names the pipeline expects.",
     None),

    ("What is the anti-hallucination guarantee?",
     "This project uses the word 'hallucination' in the AI sense: making up something that "
     "sounds plausible but is not grounded in evidence. The anti-hallucination guarantee means: "
     "every explanation or hypothesis produced by Stage 5 must be traceable back to a specific "
     "feature importance score or residual value from the actual model. No general statements "
     "about batteries that are not supported by what the model actually learned are included.",
     None),

    # ── RESULTS AND TARGETS ───────────────────────────────────────────────────
    ("RESULTS AND TARGETS", None, None),

    ("What was the main accuracy target and did we meet it?",
     "The target was: XGBoost RMSE < 100 cycles on the test set. "
     "We achieved: RMSE = 22.08 cycles. Well below the threshold. YES, target met.",
     None),

    ("What was the uncertainty target and did we meet it?",
     "Target: 90% conformal coverage — the true RUL must fall inside our predicted interval "
     "at least 90% of the time. We achieved: 98.3% overall. YES, target exceeded.",
     None),

    ("What is CONDITIONAL PASS and why did we get it instead of a full PASS?",
     "CONDITIONAL PASS means the system works well but with a noted caveat. "
     "Our caveat: the cross-validation RMSE (48.17) is 2.2 times higher than the holdout "
     "test RMSE (22.08). This means the test batteries were easier to predict than average. "
     "The system is valid for research use but needs more data to be fully reliable on all battery types.",
     None),

    ("Did we achieve everything we set out to do?",
     "Yes, all required outputs were produced: RUL predictions with uncertainty bounds, "
     "risk categories (LOW/MEDIUM/HIGH), survival risk analysis, anomaly detection, "
     "degradation hypotheses, counterfactual examples, and a supervisor-audited final report. "
     "The only two gaps were: cold-battery coverage at 88.7% (target 90%), and the CV/test "
     "RMSE ratio of 2.2x. Both are data-volume limitations, not design failures.",
     None),

    # ── GENERAL / PRESENTATION ────────────────────────────────────────────────
    ("GENERAL / PRESENTATION QUESTIONS", None, None),

    ("Can this system be used to control a real battery?",
     "No. This is a research prototype and decision-support tool. "
     "It is not validated for safety-critical use, it is not a Battery Management System (BMS), "
     "and it cannot automatically control or shut down a battery. All outputs are informational "
     "and require expert review before any action is taken.",
     None),

    ("Why not use a big neural network like GPT or a transformer instead of XGBoost?",
     "With only 34 batteries (roughly 1,700 rows of data), large neural networks would overfit — "
     "they memorise the training data rather than learning generalizable patterns. XGBoost is "
     "specifically well-suited to small tabular datasets. The TCN was a good choice for the "
     "sequence model because it has fewer parameters than a transformer and trains faster.",
     None),

    ("What would improve this project with more time/data?",
     "1. More batteries — especially more cold-temperature batteries to fix the coverage gap.\n"
     "2. More battery chemistries — NMC, LFP, not just LiCoO2.\n"
     "3. Real-world deployment data (lab conditions are more uniform than field conditions).\n"
     "4. Online learning — updating the model as new cycles arrive, without full retraining.\n"
     "5. Better deep learning architecture with more data (transformer or LSTM).",
     None),

    ("What is the knowledge base used for?",
     "The knowledge base is a folder of text files containing background information about "
     "battery degradation (e.g. descriptions of SEI layer growth, NASA/CALCE dataset details, "
     "anomaly detection methods). The supervisor report uses a TF-IDF search over these files "
     "to retrieve relevant snippets that support the system's explanations — a simple form of "
     "Retrieval-Augmented Generation (RAG).",
     None),

    ("What is TF-IDF?",
     "TF-IDF stands for Term Frequency - Inverse Document Frequency. It is a way to find "
     "text that is relevant to a query. High TF means the word appears often in a document. "
     "High IDF means the word is rare across all documents (so it is distinctive). "
     "Multiplying them gives a relevance score. It is much simpler than modern embedding-based "
     "search but works well for small knowledge bases.",
     None),

    ("What is the difference between RUL prediction and survival analysis?",
     "RUL prediction gives you a specific number: 'this battery has 47 cycles left'. "
     "Survival analysis gives you a probability: 'there is a 65% chance this battery fails "
     "within the next 20 cycles'. They answer different questions and are both useful. "
     "RUL is better for planning ('when to order a replacement'). Survival is better for "
     "risk assessment ('how urgent is this?').",
     None),

    ("What is quantile regression?",
     "Regular regression predicts the average (mean) outcome. Quantile regression predicts "
     "specific percentiles. We train three XGBoost models: one to predict the 5th percentile "
     "of RUL (pessimistic), one for the 50th percentile (median/typical), and one for the "
     "95th percentile (optimistic). The gap between the 5th and 95th percentile gives us a "
     "90% prediction interval from the model itself (before conformal calibration).",
     None),

    ("What is the difference between the quantile regression interval and the conformal interval?",
     "The quantile regression interval comes from the model itself and has no formal guarantee — "
     "if the model is wrong, the interval can be systematically too narrow. "
     "The conformal interval is post-hoc calibration using actual errors on held-out batteries. "
     "It has a mathematical guarantee regardless of model quality. "
     "We use conformal as the final interval because it is more trustworthy.",
     None),

    ("Why did we use 6 calibration batteries?",
     "This is set in the configuration file (n_cal_batteries: 6). It is a trade-off: "
     "more calibration batteries = better-calibrated q_hat but fewer training batteries. "
     "With only 28 training batteries available (80% of 34), we allocated 6 to calibration, "
     "leaving 22 for actual model training. With very few calibration batteries, a safety "
     "inflation factor of 1.20x is applied to q_hat for small groups.",
     None),
]


def build():
    print("Building Q&A document...")
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Questions and Answers")
    tr.font.size = Pt(26); tr.font.bold = True; tr.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Battery AI Co-Scientist — Complete Q&A Guide\nWritten simply for first-time learners")
    sr.font.size = Pt(13); sr.font.italic = True; sr.font.color.rgb = TEAL

    doc.add_paragraph()
    intro = doc.add_paragraph()
    r = intro.add_run(
        "This document answers every question you might be asked about this project — "
        "from the very basics to the technical details. If you are reading this for the "
        "first time, start from the top. Each answer is written to be as simple as possible. "
        "Technical terms are always explained in plain language."
    )
    r.font.size = Pt(11); r.font.color.rgb = GREY
    intro.paragraph_format.space_after = Pt(8)
    doc.add_page_break()

    # ── Table of contents (manual) ────────────────────────────────────────────
    toc_p = doc.add_paragraph()
    tr2 = toc_p.add_run("Contents")
    tr2.font.size = Pt(14); tr2.font.bold = True; tr2.font.color.rgb = NAVY
    doc.add_paragraph()

    sections_list = [
        ("Basics", "What is the project, the dataset, and key concepts"),
        ("The Pipeline / Stages", "What each of the 6 stages does"),
        ("The Models", "XGBoost, TCN, statistical baseline, RMSE, cross-validation"),
        ("Uncertainty and Confidence Intervals", "Conformal prediction, q_hat, LOBO"),
        ("Survival Analysis", "Kaplan-Meier, hazard rate, risk thresholds"),
        ("Drift Monitoring", "PSI, what RED drift means"),
        ("Anomaly Detection", "What anomalies are and how they are found"),
        ("Validation and Robustness", "Hostile tests, schema validation, anti-hallucination"),
        ("Results and Targets", "Did we achieve everything? What is CONDITIONAL PASS?"),
        ("General / Presentation Questions", "Why these choices, what would improve it, terminology"),
    ]
    for title, desc in sections_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(f"{title}  ")
        r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NAVY
        r2 = p.add_run(f"— {desc}")
        r2.font.size = Pt(11); r2.font.color.rgb = GREY

    doc.add_page_break()

    # ── Q&A ───────────────────────────────────────────────────────────────────
    q_number = 0

    for item in QA_DATA:
        q, a, note_text = item

        # Section header
        if a is None and note_text is None:
            p = doc.add_paragraph()
            r = p.add_run(q)
            r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            # Shade the paragraph background via a table trick
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.rows[0].cells[0]
            _shade(cell, "1F3A5F")
            cell.paragraphs[0].clear()
            pr = cell.paragraphs[0].add_run(q)
            pr.font.size = Pt(12); pr.font.bold = True; pr.font.color.rgb = WHITE
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
            # remove the floating paragraph above (it would duplicate)
            p._element.getparent().remove(p._element)
            doc.add_paragraph()
            continue

        q_number += 1

        # Question block
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(10)
        q_para.paragraph_format.space_after  = Pt(2)
        q_r = q_para.add_run(f"Q{q_number}.  {q}")
        q_r.font.size = Pt(11.5); q_r.font.bold = True; q_r.font.color.rgb = NAVY

        # Answer block — use a shaded table cell for visual separation
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        _shade(cell, "F0F4F8")
        cell.paragraphs[0].clear()

        # Handle multi-line answers (split on \n)
        lines = a.split("\n")
        for li, line in enumerate(lines):
            if li == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(11); run.font.color.rgb = GREY
            para.paragraph_format.space_after = Pt(2)

        # Optional note / formula
        if note_text:
            np_ = cell.add_paragraph()
            np_.paragraph_format.space_before = Pt(4)
            np_.paragraph_format.left_indent = Cm(0.5)
            nr = np_.add_run(f"Note: {note_text}")
            nr.font.size = Pt(10); nr.font.italic = True; nr.font.color.rgb = TEAL

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Final note ────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("End of Q&A Document")
    fr.font.size = Pt(14); fr.font.bold = True; fr.font.color.rgb = NAVY

    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = fp2.add_run(
        "If a question comes up during your presentation that is not covered here, "
        "the honest answer is always: 'That is a good point — the current version does not handle that, "
        "but it would be a good direction for future work.'"
    )
    fr2.font.size = Pt(11); fr2.font.italic = True; fr2.font.color.rgb = TEAL

    out = ROOT / "QA_Document.docx"
    doc.save(str(out))
    print(f"Saved: {out}")

if __name__ == "__main__":
    build()
